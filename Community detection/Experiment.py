"""Experimental driver for comparing GraphRAG community detection methods.
Based on Microsoft Research's "From Local to Global: A Graph RAG Approach to
Query-Focused Summarization" (2024).
"""
### set TF_ENABLE_ONEDNN_OPTS=0 to avoid oneDNN errors
### set OPENAI_API_KEY

import json
import logging
import os
import random
import time
from pathlib import Path
from typing import Any, Dict, List, Tuple

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import networkx as nx
import numpy as np
from openai import OpenAI

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. Install with: pip install python-docx")

CLIENT = OpenAI()
random.seed(42)
np.random.seed(42)

from microsoft_graphrag_replication import (
    MicrosoftGraphRAG, 
    LeidenCommunityDetection,
    GlobalSearchEngine,
    CommunitySummarizer,
    create_communities_dataframe,
    create_entities_dataframe,
)

from algorithms import (
    SpectralCommunityDetection,
    SCORECommunityDetection,
    NAC1CommunityDetection,
    NAC2CommunityDetection,
    clear_all_caches as clear_graphrag_caches,
)

class AdjustableLeidenCommunityDetection:
    """Leiden with adjustable max_cluster_size to control number of communities"""
    
    def __init__(self, target_k: int = None):
        self.target_k = target_k
        self._base_max_cluster_size = 50
        
    def detect_communities(self, graph: nx.Graph, **kwargs) -> List[List[str]]:
        """Detect communities using Leiden algorithm with adjustable max_cluster_size"""
        if graph.number_of_nodes() == 0:
            return []
        
        if self.target_k is None:
            return LeidenCommunityDetection().detect_communities(graph)
        
        best_communities = None
        best_diff = float('inf')
        best_size = None
        
        n_nodes = graph.number_of_nodes()
        
        # Multi-phase search to find a max_cluster_size that matches the requested k.
        # Phase 1: coarse sweep to bracket the right order of magnitude.
        coarse_sizes = []
        for divisor in [2, 3, 4, 5, 6, 8, 10, 15, 20, 30, 40, 50, 75, 100]:
            size = max(2, n_nodes // divisor)
            if size not in coarse_sizes and size < n_nodes:
                coarse_sizes.append(size)
        
        k_to_size = {}
        
        for test_size in sorted(coarse_sizes):
            try:
                from microsoft_graphrag_replication import _compute_leiden_hierarchy_local
                mapping, _parent = _compute_leiden_hierarchy_local(
                    graph=graph, max_cluster_size=test_size, use_lcc=False, seed=42
                )
                root = mapping.get(0, {})
                clusters = {}
                for node, cid in root.items():
                    clusters.setdefault(cid, []).append(node)
                
                k = len(clusters)
                k_to_size[k] = test_size
                
                diff = abs(k - self.target_k)
                if diff < best_diff:
                    best_diff = diff
                    best_communities = list(clusters.values())
                    best_size = test_size
                
                if k == self.target_k:
                    return best_communities
            except Exception as e:
                logger.warning(f"Failed with max_cluster_size={test_size}: {e}")
                continue
        
        # Phase 2: refine around the best coarse candidate.
        if best_size is not None:
            fine_sizes = []
            for offset in [-5, -3, -2, -1, 1, 2, 3, 5]:
                fine_size = best_size + offset
                if 2 <= fine_size < n_nodes and fine_size not in coarse_sizes:
                    fine_sizes.append(fine_size)
            
            for test_size in sorted(fine_sizes):
                try:
                    from microsoft_graphrag_replication import _compute_leiden_hierarchy_local
                    mapping, _parent = _compute_leiden_hierarchy_local(
                        graph=graph, max_cluster_size=test_size, use_lcc=False, seed=42
                    )
                    root = mapping.get(0, {})
                    clusters = {}
                    for node, cid in root.items():
                        clusters.setdefault(cid, []).append(node)
                    
                    k = len(clusters)
                    diff = abs(k - self.target_k)
                    if diff < best_diff:
                        best_diff = diff
                        best_communities = list(clusters.values())
                        best_size = test_size
                    
                    if k == self.target_k:
                        return best_communities
                except Exception as e:
                    logger.warning(f"Failed with max_cluster_size={test_size}: {e}")
                    continue
        
        # Phase 3: adjust by splitting or merging if we're still off by a small margin.
        if best_communities is not None and best_diff <= 2:
            if len(best_communities) < self.target_k:
                # Split largest communities
                while len(best_communities) < self.target_k:
                    largest_idx = max(range(len(best_communities)), key=lambda i: len(best_communities[i]))
                    largest = best_communities[largest_idx]
                    if len(largest) > 1:
                        mid = len(largest) // 2
                        new_comm = largest[mid:]
                        best_communities[largest_idx] = largest[:mid]
                        best_communities.append(new_comm)
                    else:
                        break
            elif len(best_communities) > self.target_k:
                # Merge smallest communities
                while len(best_communities) > self.target_k:
                    smallest_idx = min(range(len(best_communities)), key=lambda i: len(best_communities[i]))
                    second_smallest_idx = min(
                        [i for i in range(len(best_communities)) if i != smallest_idx],
                        key=lambda i: len(best_communities[i])
                    )
                    # Merge into second smallest
                    best_communities[second_smallest_idx].extend(best_communities[smallest_idx])
                    best_communities.pop(smallest_idx)
        
        if best_communities is None:
            return LeidenCommunityDetection().detect_communities(graph)
        
        return best_communities
    
    def get_name(self) -> str:
        return "adjusted_leiden"

EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "text-embedding-3-large")

# LLM choices for each stage.
EVALUATION_MODEL = "gpt-4o-mini" 

# Cache LLM judgments between runs.
JUDGE_CACHE_FILE = Path("judge_cache.json")

def _load_judge_cache() -> Dict[Tuple[str, str, str, str], str]:
    if not JUDGE_CACHE_FILE.exists():
        return {}
    try:
        with open(JUDGE_CACHE_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
        return {
            (item.get("metric", ""), item.get("query", ""), item.get("ans_i", ""), item.get("ans_j", "")):
            item.get("decision", "TIE")
            for item in data
        }
    except Exception:
        return {}

def _save_judge_cache(cache: Dict[Tuple[str, str, str, str], str]) -> None:
    try:
        serializable = [
            {"metric": k[0], "query": k[1], "ans_i": k[2], "ans_j": k[3], "decision": v}
            for k, v in cache.items()
        ]
        with open(JUDGE_CACHE_FILE, "w", encoding="utf-8") as f:
            json.dump(serializable, f, ensure_ascii=False)
    except Exception:
        # Avoid crashing if cache persistence fails.
        return

# Global in-memory cache
judge_cache: Dict[Tuple[str, str, str, str], str] = _load_judge_cache()

# Helper to adapt external algorithms to the expected interface
class AlgorithmAdapter:
    def __init__(self, name: str, impl: Any):
        self._name = name
        self._impl = impl

    def get_name(self) -> str:
        return self._name

    def detect_communities(self, graph, **kwargs):
        # Delegate to underlying implementation with robust error handling
        if hasattr(self._impl, 'detect_communities') and callable(self._impl.detect_communities):
            try:
                # Input validation
                if graph.number_of_nodes() == 0:
                    return []
                
                # Run the algorithm
                communities = self._impl.detect_communities(graph, **kwargs)
                
                # Output validation and normalization
                if not communities:
                    logger.warning(f"{self._name}: No communities returned, using fallback")
                    return self._fallback_communities(graph)
                
                # Ensure all nodes are assigned
                all_nodes = set(graph.nodes())
                assigned_nodes = {node for comm in communities for node in comm}
                
                missing_nodes = all_nodes - assigned_nodes
                if missing_nodes:
                    logger.warning(f"{self._name}: {len(missing_nodes)} nodes not assigned, adding to largest community")
                    if communities:
                        largest_idx = max(range(len(communities)), key=lambda i: len(communities[i]))
                        communities[largest_idx].extend(missing_nodes)
                
                return communities
                
            except TypeError:
                # Backward compatibility if impl doesn't accept kwargs
                try:
                    return self._impl.detect_communities(graph)
                except Exception as e:
                    logger.error(f"{self._name} failed: {e}")
                    return self._fallback_communities(graph)
            except Exception as e:
                logger.error(f"{self._name} failed: {e}")
                return self._fallback_communities(graph)
        raise AttributeError(f"Underlying implementation for {self._name} lacks detect_communities")
    
    def _fallback_communities(self, graph):
        """Fallback community detection using connected components"""
        return [list(comp) for comp in nx.connected_components(graph)]

# Document loading helpers

def load_json_documents(folder_path):
    """Load JSON documents from a folder"""
    documents = []
    titles = []
    folder = Path(folder_path)
    
    if not folder.exists():
        return documents, titles
    
    for file_path in folder.glob("*.json"):
        try:
            with open(file_path, 'r', encoding='utf-8-sig') as f:
                data = json.load(f)
                
            if isinstance(data, dict) and 'text' in data:
                text = data['text']
                title = data.get('title', file_path.stem)
            elif isinstance(data, dict) and 'content' in data:
                text = data['content']
                title = data.get('title', file_path.stem)
            else:
                text = str(data)
                title = file_path.stem
                
            if text and text.strip():
                documents.append(text)
                titles.append(title)
                
        except Exception:
            continue

    
    return documents, titles

def load_multiple_documents(folder_path=None, file_paths=None, max_documents=40):
    """Load documents from folder or specific files"""
    documents = []
    titles = []
    
    if folder_path:
        folder = Path(folder_path)
        if folder.exists():
            json_docs, json_titles = load_json_documents(folder_path)
            documents.extend(json_docs[:max_documents])
            titles.extend(json_titles[:max_documents])
    
    elif file_paths:
        for file_path in file_paths[:max_documents]:
            file_path = Path(file_path)
            if file_path.suffix.lower() == '.json':
                try:
                    with open(file_path, 'r', encoding='utf-8-sig') as f:
                        data = json.load(f)
                    
                    if isinstance(data, dict) and 'text' in data:
                        text = data['text']
                        title = data.get('title', file_path.stem)
                    else:
                        text = str(data)
                        title = file_path.stem
                        
                    documents.append(text)
                    titles.append(title)
                except Exception:
                    continue
    
    return documents, titles

# Query generation

def generate_microsoft_algorithm1_questions(documents, titles, K=3, N=3, M=3):
    """Microsoft's Algorithm 1 using explicit persona/task/question formats (no emojis).

    Args:
        documents: List of document strings
        titles: List of document titles
        K: Number of users
        N: Number of tasks per user
        M: Number of questions per (user, task)
    """

    # Build corpus description from titles and short samples of all documents
    doc_samples = []
    if documents and titles:
        per_doc_chars = 400
        doc_samples = [f"- {title}: {(doc[:per_doc_chars] + '...') if len(doc) > per_doc_chars else doc}" 
                       for doc, title in zip(documents, titles)]

    dataset_description = (
        f"A private corpus of {len(documents)} documents.\n\n"
        f"Titles: {', '.join(titles)}\n\n"
        f"Content samples (each line corresponds to one document):\n"
        f"{chr(10).join(doc_samples)}\n\n"
        "All tasks and questions must be grounded in this corpus only; do not rely on outside knowledge."
    )

    # Step 1: Generate K user personas

    personas_prompt = f"""
Based on the following corpus description, describe personas of {K} potential users of this dataset.

Corpus Description:
{dataset_description}

For each user persona, provide:
1. A clear user type/role (e.g., "Technology Journalist", "Venture Capitalist")
2. Their background and expertise
3. Why they would be interested in this dataset

Format your response as:
User 1: [Role] - [Description]
User 2: [Role] - [Description]
...
"""

    try:
        resp = CLIENT.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": personas_prompt}],
            temperature=0.3
        )
        personas_text = resp.choices[0].message.content.strip()

    except Exception as e:

        personas_text = "\n".join([f"User {i+1}: Analyst - Fallback persona" for i in range(K)])

    # Parse personas matching "User X: ..." or numbered lines
    personas: List[str] = []
    for line in personas_text.splitlines():
        s = line.strip()
        if not s:
            continue
        if ("User" in s) or (s and s[0].isdigit()):
            if ':' in s:
                persona = s.split(':', 1)[1].strip()
                if persona:
                    personas.append(persona)
    personas = personas[:K]

    # Parsed personas
    
    # Step 2: For each user, generate N tasks

    all_user_tasks: Dict[str, List[str]] = {}
    for i, persona in enumerate(personas):
        tasks_prompt = f"""
User Persona: {persona}

Corpus Description: {dataset_description}

Based on this user persona and the corpus, identify {N} simple, practical tasks that this user would want to accomplish. Each task should:
- Be relevant to the user's role
- Be brief and focused (e.g., "Understanding tech policy views", "Finding innovation trends")
- Be concrete and achievable

Format your response as:
Task 1: [Task description]
Task 2: [Task description]
...
"""

        try:
            resp = CLIENT.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": tasks_prompt}],
                temperature=0.7
            )
            tasks_text = resp.choices[0].message.content.strip()

            # Parse tasks matching "Task X: ..." or numbered
            tasks: List[str] = []
            for line in tasks_text.splitlines():
                s = line.strip()
                if not s:
                    continue
                if ("Task" in s) or (s and s[0].isdigit()):
                    if ':' in s:
                        task = s.split(':', 1)[1].strip()
                        if task:
                            tasks.append(task)

            all_user_tasks[persona] = tasks[:N]

        except Exception as e:

            all_user_tasks[persona] = []

    # Step 3: For each (user, task) pair, generate M questions

    all_questions: List[str] = []

    for user_idx, (persona, tasks) in enumerate(all_user_tasks.items()):
        for task_idx, task in enumerate(tasks):
            questions_prompt = f"""
User Persona: {persona}
User Task: {task}

Corpus Description: {dataset_description}

Generate {M} simple, practical questions that this user would ask to complete their task. 

Good question examples:
- Which episodes discuss privacy regulations?
- How do tech leaders view government oversight?
- What are the main concerns about AI ethics mentioned?
- Do any guests disagree about innovation policies?
- What successful collaborations are described?

Each question should:
- Be clear and direct
- Focus on finding specific information or themes
- Be relevant to the user's task
- Stay simple and concrete

Format your response as:
Question 1: [Question]
Question 2: [Question]
...
"""

            try:
                resp = CLIENT.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[{"role": "user", "content": questions_prompt}],
                    temperature=0.3
                )
                questions_text = resp.choices[0].message.content.strip()

                # Parse questions matching "Question X: ..." or numbered
                questions: List[str] = []
                for line in questions_text.splitlines():
                    s = line.strip()
                    if not s:
                        continue
                    if ("Question" in s) or (s and s[0].isdigit()):
                        if ':' in s:
                            question = s.split(':', 1)[1].strip()
                            if question:
                                questions.append(question)

                user_task_questions = questions[:M]
                all_questions.extend(user_task_questions)

            except Exception:
                continue
    return all_questions

# Evaluation helpers

def judge_answers(question, answer_a, answer_b, metric):
    """Judge which answer is better for a given metric with 1 evaluation (faster)."""
    
    def _normalize_for_judging(text, max_chars=1400):
        """Normalize answer length to reduce verbosity bias"""
        if len(text) <= max_chars:
            return text
        # Truncate to max_chars
        truncated = text[:max_chars]
        last_period = truncated.rfind('.')
        last_question = truncated.rfind('?')
        last_exclaim = truncated.rfind('!')
        
        last_sentence_end = max(last_period, last_question, last_exclaim)
        if last_sentence_end > max_chars * 0.8: 
            return truncated[:last_sentence_end + 1]
        else:
            return truncated + "..."
    
    # Normalize lengths to cap both answers to minimum length  
    len_a, len_b = len(answer_a), len(answer_b)
    min_len = min(len_a, len_b)
    if min_len < 1400:  # Only apply if both are reasonable length
        answer_a_norm = _normalize_for_judging(answer_a, min_len)
        answer_b_norm = _normalize_for_judging(answer_b, min_len)
    else:
        answer_a_norm = _normalize_for_judging(answer_a)
        answer_b_norm = _normalize_for_judging(answer_b)
    
    # Define metric descriptions
    metric_descriptions = {
        "Comprehensiveness": "Detail and coverage: How comprehensive is the answer in addressing the topic? Consider depth of explanation, breadth of coverage, and completeness of information",
        "Diversity": "Variety of perspectives: How diverse are the viewpoints, approaches, or insights presented? Consider multiple angles, different stakeholder perspectives, and varied analytical approaches", 
        "Empowerment": "Actionable insights: How well does the answer help informed decision-making? Consider practical applicability, clear guidance, and useful recommendations",
        "Directness": "Focus and relevance: How directly does the answer address the specific question asked? Consider relevance to the query, clarity of response, and avoidance of tangential information"
    }
    
    description = metric_descriptions.get(metric, "overall quality")
    
    base_prompt = f"""You are an expert evaluator. Judge which answer is better based on {description}.

Question: {question}

Answer A: {answer_a_norm}

Answer B: {answer_b_norm}

Decision rules:
- Prefer answers that are specific, evidence-based, and clearly grounded in the question.
- Penalize vagueness, repetition, or lack of structure.
- Prefer answers that show deeper understanding of relationships and structure in the data.
- Ties are rare; only use TIE if both answers are truly indistinguishable in quality.

Respond with exactly one word: "A", "B", or "TIE" """
    
    try:
        response = CLIENT.chat.completions.create(
            model=EVALUATION_MODEL,
            messages=[{"role": "user", "content": base_prompt}],
            temperature=0.1,
            max_tokens=1,
        )
        result = response.choices[0].message.content.strip().upper()
        return result if result in {"A", "B", "TIE"} else "TIE"
    except Exception:
        return "TIE"

# Win-rate matrix experiment

def run_winrate_matrix_experiment_fixed_k(
    documents: List[str],
    titles: List[str],
    queries: List[str],
    fixed_k: int,
    verbose: bool = True,
    baseline_gr: MicrosoftGraphRAG = None
) -> Dict[str, Dict[str, Any]]:
    """
    Compute pairwise win-rate matrices for all algorithms at a fixed k value.
    
    Methods compared:
      - Leiden: Leiden community detection
      - Spectral: Spectral clustering 
      - SCORE: Fast community detection by SCORE 
      - NAC1: Network-Adjusted Covariates Algorithm 1
      - NAC2: Network-Adjusted Covariates Algorithm 2
    """
    metrics = ["Comprehensiveness", "Diversity", "Empowerment", "Directness"]
    results = {}

    global global_leiden_graphrag
    global_leiden_answers = {}
    leiden_level_counts: Dict[str, int] = {}
    answers: Dict[str, List[str]] = {}
    algo_valid: Dict[str, bool] = {}
    # R1(a): wall-clock seconds spent in community detection per algorithm.
    detection_timings: Dict[str, float] = {}

    baseline_available = False
    
    if baseline_gr is not None:
        baseline_available = True
        if verbose:
            print("Using pre-built graph - skipping graph construction")
    else:
        try:
            temp_leiden = LeidenCommunityDetection()
            baseline_gr = MicrosoftGraphRAG(temp_leiden)
            baseline_stats = baseline_gr.index(documents, titles, max_cluster_size=50)
            baseline_available = True
        except Exception as e:
            logger.warning(f"Failed to create baseline GraphRAG: {e}")
    

    
    if baseline_gr and baseline_gr.graph:
        n_nodes = baseline_gr.graph.number_of_nodes()
        n_edges = baseline_gr.graph.number_of_edges()
        n_components = nx.number_connected_components(baseline_gr.graph)
        density = n_edges / (n_nodes * (n_nodes - 1) / 2) if n_nodes > 1 else 0
        
        if verbose:
            print(f"\nGraph statistics:")
            print(f"  Documents: {len(documents)}")
            print(f"  Nodes (entities): {n_nodes}")
            print(f"  Edges (relationships): {n_edges}")
            print(f"  Connected components: {n_components}")
            print(f"  Graph density: {density:.4f}")
        
        run_winrate_matrix_experiment_fixed_k._last_graph_stats = {
            'n_nodes': n_nodes,
            'n_edges': n_edges,
            'n_components': n_components,
            'density': density
        }
        
        if n_components > 1:
            natural_k_estimate = max(n_components, int(np.sqrt(n_nodes)))
        else:
            natural_k_estimate = max(2, int(np.sqrt(n_nodes)))

    else:
        natural_k_estimate = 10

    
    algorithms = {}
    
    target_k = fixed_k

    nac_algorithms = {
        "nac1": AlgorithmAdapter("nac_algorithm1", 
            NAC1CommunityDetection(k=target_k, use_embeddings=True)),
        "nac2": AlgorithmAdapter("nac_algorithm2", 
            NAC2CommunityDetection(k=target_k, adaptive_alpha=True)),
    }
    
    nac_community_counts = {}
    phase1_results = {}
    
    algo_valid = {}
    method_community_counts = {} 
    
    for alg_name, algorithm in nac_algorithms.items():
            try:
                if baseline_available:
                    gr = MicrosoftGraphRAG(community_algorithm=algorithm)
                    gr.graph = baseline_gr.graph.copy()
                    gr.entities = baseline_gr.entities
                    gr.relationships = baseline_gr.relationships
                    gr.text_units = baseline_gr.text_units
                    gr.documents = baseline_gr.documents
                    gr.community_summarizer = CommunitySummarizer()
                    gr.community_reports = {}

                    title_to_entity = {e.title: e for e in gr.entities}
                    for node in gr.graph.nodes():
                        if node in title_to_entity:
                            gr.graph.nodes[node]["entity_data"] = title_to_entity[node]
                    
                    embeddings = None
                    if alg_name in ['nac1', 'nac2'] and hasattr(gr, 'entities'):
                        embedding_cache_key = f"embeddings_natural"
                        if hasattr(run_winrate_matrix_experiment_fixed_k, embedding_cache_key):
                            embeddings = getattr(run_winrate_matrix_experiment_fixed_k, embedding_cache_key)

                        else:

                            try:
                                embeddings = {}
                                batch_size = 100
                                
                                for i in range(0, len(gr.entities), batch_size):
                                    batch = gr.entities[i:i+batch_size]
                                    texts = []
                                    keys = []
                                    
                                    for entity in batch:
                                        title = entity.title
                                        desc = getattr(entity, 'description', '')
                                        text = f"{title}: {desc}" if desc else title
                                        texts.append(text)
                                        keys.append(title)
                                    
                                    response = CLIENT.embeddings.create(
                                        input=texts,
                                        model=EMBEDDING_MODEL
                                    )
                                    
                                    for j, embedding in enumerate(response.data):
                                        embeddings[keys[j]] = np.array(embedding.embedding)

                                setattr(run_winrate_matrix_experiment_fixed_k, embedding_cache_key, embeddings)
                            except Exception:

                                embeddings = None
                    
                    try:
                        import signal
                        import threading
                        
                        communities_dict = None
                        error_msg = None
                        
                        def run_detection():
                            nonlocal communities_dict, error_msg
                            try:
                                if hasattr(algorithm, 'detect_communities'):
                                    communities_dict = gr.community_detector.detect_hierarchical_communities(
                                        gr.graph, 
                                        strict_k=False,
                                        min_size=2,
                                        entities=gr.entities,
                                        embeddings=embeddings
                                    )
                                else:
                                    communities_dict = gr.community_detector.detect_hierarchical_communities(gr.graph, strict_k=False, min_size=2)
                            except Exception as e:
                                error_msg = str(e)
                        
                        # R1(a): time community detection (wall clock, includes thread join).
                        _t0 = time.perf_counter()
                        thread = threading.Thread(target=run_detection)
                        thread.start()
                        thread.join(timeout=60)  # 60 second timeout
                        detection_timings[alg_name] = time.perf_counter() - _t0

                        if thread.is_alive():

                            communities_dict = {0: []}
                            nodes = list(gr.graph.nodes())
                            fallback_k = max(2, int(np.sqrt(len(nodes))))
                            for i in range(fallback_k):
                                start_idx = i * len(nodes) // fallback_k
                                end_idx = (i + 1) * len(nodes) // fallback_k if i < fallback_k - 1 else len(nodes)
                                comm_nodes = nodes[start_idx:end_idx]
                                if comm_nodes:
                                    comm = type('Community', (), {
                                        'id': f'timeout_comm_{i}',
                                        'entities': comm_nodes,
                                        'entity_ids': comm_nodes,
                                        'relationships': [],
                                        'relationship_ids': []
                                    })()
                                    communities_dict[0].append(comm)
                        elif error_msg:

                            raise Exception(error_msg)
                    except Exception as e:

                        communities_dict = {0: []}
                    if 0 in communities_dict:
                        total_communities = len(communities_dict[0])
                    else:
                        total_communities = sum(len(comms) for comms in communities_dict.values())
                    # Check for singleton communities (likely a failure mode)
                    all_sizes = []
                    for level, comms in communities_dict.items():
                        for comm in comms:
                            entities = getattr(comm, 'entities', getattr(comm, 'entity_ids', []))
                            all_sizes.append(len(entities))
                    
                    # Store initial community count before any normalization.
                    initial_communities = total_communities
                    method_community_counts[alg_name] = initial_communities

                    if communities_dict and any(communities_dict.values()):
                        # Normalise entity identifiers because algorithms may return names instead of IDs.
                        fixed_communities_dict = {}

                        # Build relationship lookups keyed by title since relationships use names.
                        entity_relationships = {}
                        entity_relationships_lower = {}

                        for rel in gr.relationships:
                            # The source and target in relationships are entity titles/names
                            key = (rel.source, rel.target)
                            if key not in entity_relationships:
                                entity_relationships[key] = []
                            entity_relationships[key].append(rel.id)
                            # Also store reverse direction
                            rev_key = (rel.target, rel.source)
                            if rev_key not in entity_relationships:
                                entity_relationships[rev_key] = []
                            entity_relationships[rev_key].append(rel.id)
                            
                            # Also store case-insensitive versions
                            key_lower = (rel.source.lower(), rel.target.lower())
                            if key_lower not in entity_relationships_lower:
                                entity_relationships_lower[key_lower] = []
                            entity_relationships_lower[key_lower].append(rel.id)
                            rev_key_lower = (rel.target.lower(), rel.source.lower())
                            if rev_key_lower not in entity_relationships_lower:
                                entity_relationships_lower[rev_key_lower] = []
                            entity_relationships_lower[rev_key_lower].append(rel.id)

                        for level, communities in communities_dict.items():
                            fixed_communities = []
                            for comm in communities:
                                # Get entity names from community
                                entity_ids = getattr(comm, 'entities', getattr(comm, 'entity_ids', []))

                                # Validate that all entity IDs exist in the graph.
                                valid_entity_ids = []
                                invalid_ids = []
                                
                                for entity_id in entity_ids:
                                    if entity_id in gr.graph.nodes():
                                        valid_entity_ids.append(entity_id)
                                    else:
                                        invalid_ids.append(entity_id)
                                        logger.debug(f"Invalid entity ID not in graph: {entity_id}")
                                
                                if invalid_ids and len(invalid_ids) < 5:
                                    logger.warning(f"Invalid entity IDs in {alg_name}: {invalid_ids}")
                                
                                # Keep only valid entity IDs.
                                entity_ids = valid_entity_ids

                                # Gather relationship IDs associated with the resolved entities.
                                relationship_ids = []
                                for i, entity_id1 in enumerate(valid_entity_ids):
                                    for j, entity_id2 in enumerate(valid_entity_ids[i+1:], i+1):
                                            title1 = gr.graph.nodes[entity_id1].get('title', entity_id1)
                                            title2 = gr.graph.nodes[entity_id2].get('title', entity_id2)
                                            
                                            key = (title1, title2)
                                            if key in entity_relationships:
                                                relationship_ids.extend(entity_relationships[key])
                                                logger.debug(f"Found relationship {title1} -> {title2}: {entity_relationships[key]}")
                                            else:
                                                # Check reverse direction
                                                rev_key = (title2, title1)
                                                if rev_key in entity_relationships:
                                                    relationship_ids.extend(entity_relationships[rev_key])
                                                    logger.debug(f"Found relationship {title2} -> {title1}: {entity_relationships[rev_key]}")
                                                else:
                                                    # Try case-insensitive matching for relationship lookup
                                                    key_lower = (title1.lower(), title2.lower())
                                                    if key_lower in entity_relationships_lower:
                                                        relationship_ids.extend(entity_relationships_lower[key_lower])
                                                        logger.debug(f"Found relationship via case-insensitive match {title1}-{title2}")
                                                    else:
                                                        logger.debug(f"No relationship found between {title1} and {title2}")

                                # Remove duplicates
                                relationship_ids = list(set(relationship_ids))

                                # Enforce minimal viability: at least 2 entities
                                # Allow disconnected components (no internal edges) for consistency
                                if len(entity_ids) < 2:
                                    logger.info(f"Skipping {alg_name} L{level}C{i}: <2 resolved entities")
                                    continue

                                # Don't skip communities with no internal edges - they're valid disconnected components
                                if not relationship_ids:
                                    logger.info(f"Note: {alg_name} L{level}C{i} has no internal edges (disconnected component)")
                                    # Continue processing instead of skipping

                                # Update community with resolved entity and relationship IDs.
                                if hasattr(comm, 'entity_ids'):
                                    comm.entity_ids = entity_ids
                                elif hasattr(comm, 'entities'):
                                    comm.entities = entity_ids
                                    comm.entity_ids = entity_ids

                                if hasattr(comm, 'relationship_ids'):
                                    comm.relationship_ids = relationship_ids
                                elif hasattr(comm, 'relationships'):
                                    comm.relationships = relationship_ids

                                fixed_communities.append(comm)

                            fixed_communities_dict[level] = fixed_communities

                        if not any(fixed_communities_dict.values()):
                            method_community_counts[alg_name] = initial_communities
                        else:
                            communities_dict = fixed_communities_dict
                            gr.communities = fixed_communities_dict
                            total_communities_after_fix = sum(len(comms) for comms in fixed_communities_dict.values())
                            try:
                                gr.community_reports = gr.community_summarizer.generate_community_reports(
                                    communities_dict, gr.graph
                                )

                                communities_df = create_communities_dataframe(communities_dict)
                                entities_df = create_entities_dataframe(gr.entities, communities_dict)
                                gr.search_engine = GlobalSearchEngine(
                                    gr.community_reports,
                                    index_communities_df=communities_df,
                                    index_entities_df=entities_df,
                                )
                                gr.communities = communities_dict
                                algo_valid[alg_name] = True

                            except Exception:
                                gr.community_reports = {}
                                algo_valid[alg_name] = False
                    else:

                        gr.community_reports = {}
                        algo_valid[alg_name] = False
                        method_community_counts[alg_name] = 0
                    
                    if not hasattr(gr, 'community_reports') or gr.community_reports is None:
                        gr.community_reports = {}
                    
                    if not hasattr(gr, 'search_engine') or gr.search_engine is None:
                        if hasattr(gr, 'communities') and gr.communities:
                            communities_df = create_communities_dataframe(gr.communities)
                            entities_df = create_entities_dataframe(gr.entities, gr.communities)
                            gr.search_engine = GlobalSearchEngine(
                                gr.community_reports,
                                index_communities_df=communities_df,
                                index_entities_df=entities_df
                            )
                        else:
                            gr.search_engine = GlobalSearchEngine(gr.community_reports)
                    
                    alg_answers = []
                    for query in queries:
                        try:
                            answer = gr.query(query, level=0) if gr.community_reports else "No communities found"
                            alg_answers.append(answer)
                        except Exception:
                            alg_answers.append("No communities found")
                    
                    answers[alg_name] = alg_answers
                    algo_valid[alg_name] = True
                else:
                    # Fallback: try to create from scratch
                    graphrag = MicrosoftGraphRAG(algorithm)
                    graphrag.index(documents, titles, max_cluster_size=20)
                    
                    alg_answers = []
                    for query in queries:
                        try:
                            answer = graphrag.query(query, level=0)
                            alg_answers.append(answer)
                        except Exception as e:
                            alg_answers.append(f"{alg_name} failed")
                    
                    answers[alg_name] = alg_answers
                    algo_valid[alg_name] = True
                
            except Exception as e:

                answers[alg_name] = [f"{alg_name} failed"] * len(queries)
                algo_valid[alg_name] = False
                # Ensure count is set even on failure
                if alg_name not in method_community_counts:
                    method_community_counts[alg_name] = 0
    
    phase1_results = answers
    nac_community_counts = method_community_counts.copy()

    spectral_score_algorithms = {
        "spectral": AlgorithmAdapter("spectral_clustering", 
            SpectralCommunityDetection(k=target_k)),  # Use target K
        "score": AlgorithmAdapter("score", 
            SCORECommunityDetection(k=target_k)),  # Use target K
    }

    nac_with_target = {}

    for alg_name, algorithm in spectral_score_algorithms.items():

        try:
            if baseline_available:
                # Reuse baseline extraction
                gr = MicrosoftGraphRAG(community_algorithm=algorithm)
                gr.graph = baseline_gr.graph.copy()
                gr.entities = baseline_gr.entities
                gr.relationships = baseline_gr.relationships
                gr.text_units = baseline_gr.text_units
                gr.documents = baseline_gr.documents
                gr.community_summarizer = CommunitySummarizer()
                gr.community_reports = {}
                
                # Set entity data in graph nodes
                title_to_entity = {e.title: e for e in gr.entities}
                for node in gr.graph.nodes():
                    if node in title_to_entity:
                        gr.graph.nodes[node]["entity_data"] = title_to_entity[node]
                
                # Detect communities (no embeddings needed for Spectral/SCORE)

                # R1(a): time community detection (wall clock).
                _t0 = time.perf_counter()
                communities_dict = gr.community_detector.detect_hierarchical_communities(
                    gr.graph, strict_k=True, min_size=2
                )
                detection_timings[alg_name] = time.perf_counter() - _t0
                
                # For non-hierarchical algorithms, only count level 0
                if 0 in communities_dict:
                    total_communities = len(communities_dict[0])
                else:
                    total_communities = sum(len(comms) for comms in communities_dict.values())
                
                # Store initial count before filtering
                initial_communities = total_communities
                method_community_counts[alg_name] = initial_communities

                
                # Fix community entity IDs (same logic as NAC)
                if communities_dict and any(communities_dict.values()):
                    fixed_communities_dict = {}
                    entity_title_to_id = {e.title: e.id for e in gr.entities}
                    entity_relationships = {}
                    
                    for rel in gr.relationships:
                        key = (rel.source, rel.target)
                        if key not in entity_relationships:
                            entity_relationships[key] = []
                        entity_relationships[key].append(rel.id)
                        rev_key = (rel.target, rel.source)
                        if rev_key not in entity_relationships:
                            entity_relationships[rev_key] = []
                        entity_relationships[rev_key].append(rel.id)
                    
                    for level, communities in communities_dict.items():
                        fixed_communities = []
                        for comm in communities:
                            entity_ids = getattr(comm, 'entities', getattr(comm, 'entity_ids', []))
                            valid_entity_ids = [eid for eid in entity_ids if eid in gr.graph.nodes()]
                            
                            # Get relationships
                            relationship_ids = []
                            for i, entity_id1 in enumerate(valid_entity_ids):
                                for j, entity_id2 in enumerate(valid_entity_ids[i+1:], i+1):
                                    title1 = gr.graph.nodes[entity_id1].get('title', entity_id1)
                                    title2 = gr.graph.nodes[entity_id2].get('title', entity_id2)
                                    key = (title1, title2)
                                    if key in entity_relationships:
                                        relationship_ids.extend(entity_relationships[key])
                                    else:
                                        rev_key = (title2, title1)
                                        if rev_key in entity_relationships:
                                            relationship_ids.extend(entity_relationships[rev_key])
                            
                            relationship_ids = list(set(relationship_ids))
                            
                            # Allow disconnected components (no relationship requirement)
                            if len(valid_entity_ids) >= 2:
                                if hasattr(comm, 'entity_ids'):
                                    comm.entity_ids = valid_entity_ids
                                elif hasattr(comm, 'entities'):
                                    comm.entities = valid_entity_ids
                                    # Also set entity_ids for consistency in report generation
                                    comm.entity_ids = valid_entity_ids
                                if hasattr(comm, 'relationship_ids'):
                                    comm.relationship_ids = relationship_ids
                                elif hasattr(comm, 'relationships'):
                                    comm.relationships = relationship_ids
                                fixed_communities.append(comm)
                        
                        fixed_communities_dict[level] = fixed_communities
                    
                    gr.communities = fixed_communities_dict
                    
                    # Generate community reports
                    gr.community_reports = gr.community_summarizer.generate_community_reports(
                        fixed_communities_dict, gr.graph
                    )
                    
                    # Create search engine
                    communities_df = create_communities_dataframe(fixed_communities_dict)
                    entities_df = create_entities_dataframe(gr.entities, fixed_communities_dict)
                    gr.search_engine = GlobalSearchEngine(
                        gr.community_reports,
                        index_communities_df=communities_df,
                        index_entities_df=entities_df
                    )
                
                # Generate answers for all queries
                alg_answers = []
                for query in queries:
                    try:
                        if gr.community_reports:
                            answer = gr.query(query, level=0)
                        else:
                            answer = "No communities found"
                        alg_answers.append(answer)
                    except Exception as e:
                        alg_answers.append(f"{alg_name} failed")
                
                answers[alg_name] = alg_answers
                algo_valid[alg_name] = True
                
        except Exception as e:

            answers[alg_name] = [f"{alg_name} failed"] * len(queries)
            algo_valid[alg_name] = False
            method_community_counts[alg_name] = 0
    
    # Reuse NAC answers gathered in phase 1.
    for alg_name in ['nac1', 'nac2']:
        if alg_name in phase1_results:
            answers[alg_name] = phase1_results[alg_name]

    # Phase 4: run Leiden with the requested k.

    try:
        # Use custom adjustable Leiden with fixed k
        adjusted_leiden = AdjustableLeidenCommunityDetection(target_k=target_k)
        global_leiden_graphrag = MicrosoftGraphRAG(adjusted_leiden)
        
        if baseline_available and baseline_gr:
            # Reuse baseline extraction
            global_leiden_graphrag.graph = baseline_gr.graph.copy()
            global_leiden_graphrag.entities = baseline_gr.entities
            global_leiden_graphrag.relationships = baseline_gr.relationships
            global_leiden_graphrag.text_units = baseline_gr.text_units
            global_leiden_graphrag.documents = baseline_gr.documents
            global_leiden_graphrag.community_summarizer = CommunitySummarizer()
            
            # Set entity data in graph nodes
            title_to_entity = {e.title: e for e in global_leiden_graphrag.entities}
            for node in global_leiden_graphrag.graph.nodes():
                if node in title_to_entity:
                    global_leiden_graphrag.graph.nodes[node]["entity_data"] = title_to_entity[node]
            
            # Detect communities using the hierarchical detector (like other algorithms).
            # R1(a): time community detection (wall clock).
            _t0 = time.perf_counter()
            communities_dict = global_leiden_graphrag.community_detector.detect_hierarchical_communities(
                global_leiden_graphrag.graph, strict_k=False, min_size=2
            )
            detection_timings["c0"] = time.perf_counter() - _t0
            
            # Debug: confirm single level detection
            if 0 in communities_dict:
                total_communities = len(communities_dict[0])
            else:
                total_communities = sum(len(comms) for comms in communities_dict.values())
            
            # Store initial count before filtering
            initial_communities = total_communities
            
            # Fix community entity IDs (same logic as for other algorithms)
            fixed_communities_dict = {}
            entity_title_to_id = {e.title: e.id for e in global_leiden_graphrag.entities}
            entity_relationships = {}
            for rel in global_leiden_graphrag.relationships:
                key = (rel.source, rel.target)
                if key not in entity_relationships:
                    entity_relationships[key] = []
                entity_relationships[key].append(rel.id)
                rev_key = (rel.target, rel.source)
                if rev_key not in entity_relationships:
                    entity_relationships[rev_key] = []
                entity_relationships[rev_key].append(rel.id)
            
            for level, communities in communities_dict.items():
                fixed_communities = []
                for comm in communities:
                    entity_ids = getattr(comm, 'entities', getattr(comm, 'entity_ids', []))
                    valid_entity_ids = [eid for eid in entity_ids if eid in global_leiden_graphrag.graph.nodes()]
                    
                    # Get relationships
                    relationship_ids = []
                    for i, entity_id1 in enumerate(valid_entity_ids):
                        for j, entity_id2 in enumerate(valid_entity_ids[i+1:], i+1):
                            title1 = global_leiden_graphrag.graph.nodes[entity_id1].get('title', entity_id1)
                            title2 = global_leiden_graphrag.graph.nodes[entity_id2].get('title', entity_id2)
                            key = (title1, title2)
                            if key in entity_relationships:
                                relationship_ids.extend(entity_relationships[key])
                            else:
                                rev_key = (title2, title1)
                                if rev_key in entity_relationships:
                                    relationship_ids.extend(entity_relationships[rev_key])
                    
                    relationship_ids = list(set(relationship_ids))
                    
                    # Allow disconnected components (no relationship requirement)
                    if len(valid_entity_ids) >= 2:
                        if hasattr(comm, 'entity_ids'):
                            comm.entity_ids = valid_entity_ids
                        elif hasattr(comm, 'entities'):
                            comm.entities = valid_entity_ids
                            # Also set entity_ids for consistency in report generation
                            comm.entity_ids = valid_entity_ids
                        if hasattr(comm, 'relationship_ids'):
                            comm.relationship_ids = relationship_ids
                        elif hasattr(comm, 'relationships'):
                            comm.relationships = relationship_ids
                        fixed_communities.append(comm)
                
                fixed_communities_dict[level] = fixed_communities
            
            
            global_leiden_graphrag.communities = fixed_communities_dict
            
            # Generate reports
            global_leiden_graphrag.community_reports = global_leiden_graphrag.community_summarizer.generate_community_reports(
                fixed_communities_dict, global_leiden_graphrag.graph
            )
            
            # Create search engine
            communities_df = create_communities_dataframe(fixed_communities_dict)
            entities_df = create_entities_dataframe(global_leiden_graphrag.entities, fixed_communities_dict)
            global_leiden_graphrag.search_engine = GlobalSearchEngine(
                global_leiden_graphrag.community_reports,
                index_communities_df=communities_df,
                index_entities_df=entities_df
            )
        else:
            # Create from scratch
            leiden_stats = global_leiden_graphrag.index(documents, titles, max_cluster_size=50)
        
        # Get answers for level 0 only (C0)
        level_answers = []
        for query in queries:
            try:
                answer = global_leiden_graphrag.query(query, level=0)
                level_answers.append(answer)
            except Exception as e:

                level_answers.append(f"Leiden C0 failed")
        
        global_leiden_answers["c0"] = level_answers
        
        # Use initial count before filtering
        leiden_level_counts["c0"] = initial_communities if 'initial_communities' in locals() else 0

        algo_valid["c0"] = True
        
    except Exception as e:

        global_leiden_graphrag = None
        global_leiden_answers["c0"] = ["Leiden failed"] * len(queries)
        leiden_level_counts["c0"] = 0
        algo_valid["c0"] = False
    
    answers["c0"] = global_leiden_answers.get("c0", ["Leiden failed"] * len(queries))
    all_counts = [
        ('NAC1', method_community_counts.get('nac1', 0)),
        ('NAC2', method_community_counts.get('nac2', 0)),
        ('Spectral', method_community_counts.get('spectral', 0)),
        ('SCORE', method_community_counts.get('score', 0)),
        ('Leiden C0', leiden_level_counts.get('c0', 0))
    ]

    # Labels map algorithm codes to display names.
    cat_to_label = {
        "c0": f"Leiden (k={leiden_level_counts.get('c0', target_k)})",
        "spectral": f"Spectral (k={method_community_counts.get('spectral', target_k)})",
        "score": f"SCORE (k={method_community_counts.get('score', target_k)})",
        "nac1": f"NAC1 (k={method_community_counts.get('nac1', target_k)})",
        "nac2": f"NAC2 (k={method_community_counts.get('nac2', target_k)})"
    }
    
    # Determine available methods
    available = []
    if "c0" in answers:
        available.append("c0")
    for method in ["spectral", "score", "nac1", "nac2"]:
        if algo_valid.get(method, False):
            available.append(method)

    # Function to check if answer is invalid
    def _bad(x):
        if not x or not isinstance(x, str):
            return True
        x_lower = x.lower().strip()
        markers = [
            "no communities found", "failed", "error", "not available",
            "unable to", "cannot", "empty", "null"
        ]
        return any(marker in x_lower for marker in markers)
        
    # Compute win-rate matrices for each metric 
    result = {"labels": [cat_to_label[c] for c in available]}
    n = len(available)
    
    for metric in metrics:
        matrix = [[0.0] * n for _ in range(n)]
        
        # Fill diagonal with 0.5
        for d in range(n):
            matrix[d][d] = 0.5
        
        # Evaluate unordered pairs once for deterministic symmetry
        for i in range(n):
            for j in range(i + 1, n):
                    cat_i = available[i]
                    cat_j = available[j]

                    valid_comparisons = 0
                    wins_i = 0.0
                    ties = 0.0
                    
                    for idx, query in enumerate(queries):
                        ans_i = answers[cat_i][idx]
                        ans_j = answers[cat_j][idx]
                        
                        # Skip if either answer is invalid
                        if _bad(ans_i) or _bad(ans_j):
                            continue
                        
                        # Judge this comparison with memoization to keep results stable across k
                        cache_key = (metric, query, ans_i, ans_j)
                        if cache_key in judge_cache:
                            judgment = judge_cache[cache_key]
                        else:
                            try:
                                judgment = judge_answers(query, ans_i, ans_j, metric)
                                if judgment not in ("A", "B", "TIE"):
                                    judgment = "TIE"
                            except Exception:
                                judgment = "TIE"
                            judge_cache[cache_key] = judgment
                        valid_comparisons += 1
                        
                        if judgment == "A":
                            wins_i += 1.0
                        elif judgment == "TIE":
                            ties += 1.0
                    
                    if valid_comparisons > 0:
                        rate_i = (wins_i + 0.5 * ties) / valid_comparisons
                        rate_j = 1.0 - rate_i
                        
                        if rate_i > 0.8 or rate_i < 0.2:
                            logger.info(f"Large win-rate difference: {cat_to_label[cat_i]} vs {cat_to_label[cat_j]}: {rate_i:.2f}")
                            logger.info(f"  Wins: {wins_i}/{valid_comparisons}, Ties: {ties}")
                    else:
                        rate_i = rate_j = 0.5
                        logger.warning(f"No valid comparisons between {cat_i} and {cat_j}")
                    
                    matrix[i][j] = rate_i
                    matrix[j][i] = rate_j
        
        result[metric] = matrix
    

    clear_graphrag_caches()
    
    for method in available:
        if method == "c0":
            result[f"{method}_actual_k"] = leiden_level_counts.get('c0', 0)
        else:
            result[f"{method}_actual_k"] = method_community_counts.get(method, 0)
    
    # Store the target K for graph plotting
    result["consensus_k"] = target_k
    # R1(a): expose per-algorithm community-detection wall-clock seconds.
    result["detection_timings"] = dict(detection_timings)
    if verbose and detection_timings:
        print("\nCommunity detection wall-clock (seconds):")
        for _m, _t in detection_timings.items():
            print(f"  {_m:>8}: {_t:7.3f}s")
    return {"fixed_k": result}

def compute_elo_ratings(result: Dict[str, Any], methods: List[str], K: int = 32) -> Dict[str, Dict[str, float]]:
    """Compute Elo ratings from pairwise win-rate matrices
    
    Based on the Elo rating system (Elo, 1978): "The Rating of Chessplayers, Past and Present"
    
    The Elo rating formula:
    - Expected score: E_A = 1 / (1 + 10^((R_B - R_A) / 400))
    - Rating update: R'_A = R_A + K * (S_A - E_A)
    
    Where:
    - R_A, R_B are current ratings
    - S_A is actual score (win rate in our case)
    - K is the K-factor controlling update magnitude
    
    We simulate multiple rounds of pairwise comparisons using the empirical
    win rates to converge to stable ratings.
    
    Args:
        result: Dict containing win-rate matrices for each metric
        methods: List of method names
        K: Elo K-factor (default 32, typical values: 16 for masters, 32 for regular)
    
    Returns:
        Dict mapping metric -> method -> Elo rating
    """
    metrics = ["Comprehensiveness", "Diversity", "Empowerment", "Directness"]
    elo_results = {}
    
    for metric in metrics:
        if metric not in result:
            continue
            
        matrix = result[metric]
        n = len(methods)
        
        # Initialize Elo ratings at 1500
        ratings = dict.fromkeys(methods, 1500.0)
        num_rounds = 100
        
        for _ in range(num_rounds):
            for i in range(n):
                for j in range(n):
                    if i == j:
                        continue
                    
                    method_i = methods[i]
                    method_j = methods[j]
                    
                    # Win rate of i against j
                    win_rate = matrix[i][j]
                    
                    # Expected score for i
                    expected_i = 1 / (1 + 10**((ratings[method_j] - ratings[method_i]) / 400))
                    
                    # Actual score (win rate)
                    actual_i = win_rate
                    
                    # Update ratings
                    ratings[method_i] += K * (actual_i - expected_i)
                    ratings[method_j] += K * ((1 - actual_i) - (1 - expected_i))
        
        # Normalize ratings to have mean 1500
        mean_rating = sum(ratings.values()) / len(ratings)
        for method in ratings:
            ratings[method] += (1500 - mean_rating)
        
        elo_results[metric] = ratings
    
    return elo_results

def save_elo_to_word(elo_ratings: Dict[str, Dict[str, float]], cat_to_label: Dict[str, str], 
                     k_value: int, doc: Document = None) -> Document:
    """Save ELO ratings to a Word document"""
    if not DOCX_AVAILABLE:
        logger.warning("Cannot save to Word: python-docx not installed")
        return None
    
    if doc is None:
        doc = Document()
    
    doc.add_heading(f'ELO Ratings Summary - K={k_value}', level=1)
    
    if not elo_ratings:
        doc.add_paragraph("No ELO ratings available")
        return doc
    
    metrics = list(elo_ratings.keys())
    methods = list(next(iter(elo_ratings.values())).keys())
    
    avg_elo = {}
    for method in methods:
        ratings = [elo_ratings[metric][method] for metric in metrics if method in elo_ratings[metric]]
        avg_elo[method] = sum(ratings) / len(ratings) if ratings else 0
    
    sorted_methods = sorted(methods, key=lambda m: avg_elo[m], reverse=True)
    
    # Create table
    table = doc.add_table(rows=1, cols=len(metrics) + 2)
    table.style = 'Light Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Method'
    for i, metric in enumerate(metrics):
        hdr_cells[i + 1].text = metric[:12]
    hdr_cells[len(metrics) + 1].text = 'Average'
    
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    for method in sorted_methods:
        row_cells = table.add_row().cells
        row_cells[0].text = cat_to_label.get(method, method)
        for i, metric in enumerate(metrics):
            if metric in elo_ratings and method in elo_ratings[metric]:
                row_cells[i + 1].text = f"{elo_ratings[metric][method]:.0f}"
            else:
                row_cells[i + 1].text = "N/A"
        row_cells[len(metrics) + 1].text = f"{avg_elo[method]:.0f}"
    
    # Add interpretation
    doc.add_paragraph()
    doc.add_paragraph("Interpretation:")
    doc.add_paragraph("• Higher Elo rating = better performance", style='List Bullet')
    doc.add_paragraph("• Ratings are relative within each metric", style='List Bullet')
    doc.add_paragraph("• Average provides overall ranking across all metrics", style='List Bullet')
    
    doc.add_page_break()
    
    return doc

def create_elo_progression_graphs(
    all_elo_results: Dict[int, Dict[str, Dict[str, float]]],
    methods_order: List[str] = None,
    save_plots: bool = True,
) -> None:
    """R2.M2: plot ELO rating as a function of k (= number of communities)
    using the k-values actually tested in this experiment (e.g. 10, 50, 100).

    Produces two figures:
      - elo_progression_by_k.png: one panel per metric (4 panels) with one line per method.
      - elo_progression_avg_by_k.png: a single panel of methods, with ELO averaged
        across the 4 metrics for each k.
    """
    if not all_elo_results:
        print("ELO progression skipped: no ELO results collected.")
        return

    if methods_order is None:
        methods_order = ["c0", "spectral", "score", "nac1", "nac2"]
    label_map = {"c0": "Leiden", "spectral": "Spectral", "score": "SCORE", "nac1": "NAC1", "nac2": "NAC2"}
    color_map = {"c0": "#1f77b4", "spectral": "#ff7f0e", "score": "#2ca02c",
                 "nac1": "#d62728", "nac2": "#9467bd"}
    metrics = ["Comprehensiveness", "Diversity", "Empowerment", "Directness"]

    k_values = sorted(all_elo_results.keys())

    # --- Per-metric panel (2x2) ---
    try:
        fig, axes = plt.subplots(2, 2, figsize=(14, 12))
        for idx, metric in enumerate(metrics):
            row = idx // 2
            col = idx % 2
            ax = axes[row, col]
            for method in methods_order:
                ys = []
                xs = []
                for k in k_values:
                    elo_for_metric = all_elo_results[k].get(metric, {})
                    if method in elo_for_metric:
                        xs.append(k)
                        ys.append(elo_for_metric[method])
                if ys:
                    ax.plot(xs, ys, marker='o', linewidth=2.5, markersize=10,
                            label=label_map[method], color=color_map[method])
            ax.set_title(metric, fontweight='bold', fontsize=18)
            ax.set_xlabel("Target k (number of communities)", fontsize=14)
            ax.set_ylabel("ELO rating", fontsize=14)
            ax.set_xticks(k_values)
            ax.grid(True, alpha=0.3)
            if idx == 0:
                ax.legend(loc='best', fontsize=12)
        fig.suptitle("ELO vs number of communities k", fontsize=20, fontweight='bold')
        plt.tight_layout()
        if save_plots:
            plt.savefig(os.path.join("Output", "elo_progression_by_k.png"), dpi=400, bbox_inches='tight')
            print("Saved ELO progression (per-metric): elo_progression_by_k.png")
        plt.close(fig)
    except Exception as e:
        print(f"Per-metric ELO progression plot failed: {e}")

    # --- Average ELO across metrics ---
    try:
        fig, ax = plt.subplots(figsize=(9, 6))
        for method in methods_order:
            xs, ys = [], []
            for k in k_values:
                ratings = [all_elo_results[k][m][method]
                           for m in metrics
                           if m in all_elo_results[k] and method in all_elo_results[k][m]]
                if ratings:
                    xs.append(k)
                    ys.append(sum(ratings) / len(ratings))
            if ys:
                ax.plot(xs, ys, marker='o', linewidth=2.5, markersize=11,
                        label=label_map[method], color=color_map[method])
        ax.set_title("Average ELO across 4 metrics vs k", fontsize=18, fontweight='bold')
        ax.set_xlabel("Target k (number of communities)", fontsize=14)
        ax.set_ylabel("Average ELO rating (4 metrics)", fontsize=14)
        ax.set_xticks(k_values)
        ax.grid(True, alpha=0.3)
        ax.legend(loc='best', fontsize=12)
        plt.tight_layout()
        if save_plots:
            plt.savefig(os.path.join("Output", "elo_progression_avg_by_k.png"), dpi=400, bbox_inches='tight')
            print("Saved ELO progression (average): elo_progression_avg_by_k.png")
        plt.close(fig)
    except Exception as e:
        print(f"Average ELO progression plot failed: {e}")


def run_scalability_experiment(
    baseline_gr: "MicrosoftGraphRAG",
    node_sizes: List[int] = None,
    target_k: int = 50,
    nac_timeout_seconds: float = 300.0,
    save_plot: bool = True,
    save_csv: bool = True,
    nac1_use_embeddings: bool = True,
    entity_embeddings: Dict[str, "np.ndarray"] = None,
    plot_filename: str = "scalability_curve.png",
    csv_filename: str = "scalability_timings.csv",
) -> Dict[str, Any]:
    """R1(b): scalability curve.

    Take the already-built KG, node-induce subgraphs of increasing sizes, time
    each community-detection algorithm on each subgraph, and emit a wall-clock
    vs N plot. By default, NAC1 is run *with* embeddings so its timings match
    the main experiment. Pass nac1_use_embeddings=False (and no entity_embeddings)
    to measure community-detection-only cost.

    Args:
        baseline_gr: MicrosoftGraphRAG containing .graph and (if embeddings on) .entities.
        node_sizes: subsample ladder; defaults to [50, 100, 200, 350, 500, 750, 1000, 1500].
        target_k: per-subgraph target k.
        nac_timeout_seconds: per-call timeout for NAC algorithms.
        save_plot / save_csv: write outputs.
        nac1_use_embeddings: if True (default), NAC1 receives the entity_embeddings
            so the reported wall-clock includes its embedding-aware step. Must be
            paired with a non-empty entity_embeddings dict for this to take effect.
        entity_embeddings: dict {entity_title: np.ndarray}. Required when
            nac1_use_embeddings=True; ignored otherwise.

    Returns a dict with sizes, timings, target_k.
    """
    import threading

    if baseline_gr is None or baseline_gr.graph is None:
        print("Scalability experiment skipped: no baseline graph.")
        return {}

    full_graph: nx.Graph = baseline_gr.graph
    full_nodes = list(full_graph.nodes())
    total_n = len(full_nodes)
    if total_n < 20:
        print(f"Scalability experiment skipped: graph too small (N={total_n}).")
        return {}

    if node_sizes is None:
        # Default ladder; cap at the available number of nodes.
        candidate = [50, 100, 200, 350, 500, 750, 1000, 1500]
        node_sizes = sorted({min(total_n, s) for s in candidate if s <= total_n} | {total_n})

    # Stable, reproducible subsampling: shuffle once with a fixed seed, then take prefixes.
    rng = random.Random(42)
    shuffled = list(full_nodes)
    rng.shuffle(shuffled)

    # If we're meant to use embeddings but none were passed, fall back to the
    # community-detection-only NAC1 run and warn loudly.
    effective_use_emb = bool(nac1_use_embeddings and entity_embeddings)
    if nac1_use_embeddings and not entity_embeddings:
        print("WARNING: nac1_use_embeddings=True but no entity_embeddings dict passed; "
              "falling back to community-detection-only NAC1 timing.")

    # Subset the entity list and embedding dict per subgraph so NAC1 sees only the
    # entities that exist in the subsampled subgraph.
    title_to_entity = {}
    if effective_use_emb and getattr(baseline_gr, "entities", None):
        title_to_entity = {e.title: e for e in baseline_gr.entities}

    # Algorithms to time. Spectral / SCORE need a target k; NACs need k.
    def _make_algorithms(k: int):
        return {
            "c0":       AdjustableLeidenCommunityDetection(target_k=k),
            "spectral": AlgorithmAdapter("spectral_clustering", SpectralCommunityDetection(k=k)),
            "score":    AlgorithmAdapter("score", SCORECommunityDetection(k=k)),
            "nac1":     AlgorithmAdapter("nac_algorithm1",
                                         NAC1CommunityDetection(k=k, use_embeddings=effective_use_emb)),
            "nac2":     AlgorithmAdapter("nac_algorithm2", NAC2CommunityDetection(k=k, adaptive_alpha=True)),
        }

    timings: Dict[str, Dict[int, float]] = {m: {} for m in ["c0", "spectral", "score", "nac1", "nac2"]}

    print("\n" + "=" * 80)
    print("SCALABILITY EXPERIMENT (R1.b)")
    print("=" * 80)
    print(f"Full graph: N={total_n} nodes, M={full_graph.number_of_edges()} edges")
    print(f"Subsample sizes: {node_sizes}")
    print(f"Per-subgraph target k = {target_k}")
    print(f"NAC1 embeddings: {'ON' if effective_use_emb else 'OFF'}")

    # Warmup pass: each algorithm runs once on a tiny throwaway subgraph so BLAS / LAPACK /
    # numpy / scipy initialization cost is paid here, NOT on the first real measurement at
    # the smallest N. Discard the timings.
    warmup_size = min(40, total_n)
    warmup_nodes = shuffled[:warmup_size]
    warmup_subgraph = full_graph.subgraph(warmup_nodes).copy()
    warmup_k = max(2, min(8, warmup_size // 4))
    warmup_algorithms = _make_algorithms(warmup_k)
    print(f"\nWarmup pass on N={warmup_size}, k={warmup_k} (timings discarded)...")
    for alg_name, algorithm in warmup_algorithms.items():
        warmup_kwargs = {}
        if alg_name == "nac1" and effective_use_emb:
            warmup_kwargs = {
                "entities": [title_to_entity[n] for n in warmup_nodes if n in title_to_entity],
                "embeddings": {n: entity_embeddings[n]
                               for n in warmup_nodes if n in entity_embeddings},
            }
        try:
            algorithm.detect_communities(warmup_subgraph, **warmup_kwargs)
        except Exception:
            # Warmup failures are non-fatal — the real loop has its own error handling.
            pass

    for N in node_sizes:
        # Node-induced subgraph on a deterministic prefix of the shuffled nodes.
        sub_nodes = shuffled[:N]
        subgraph = full_graph.subgraph(sub_nodes).copy()
        print(f"\n-- N={N}: M={subgraph.number_of_edges()} edges --")

        # Per-subgraph entity + embedding subsets for NAC1.
        sub_entities = None
        sub_embeddings = None
        if effective_use_emb:
            sub_entities = [title_to_entity[n] for n in sub_nodes if n in title_to_entity]
            sub_embeddings = {n: entity_embeddings[n]
                              for n in sub_nodes if n in entity_embeddings}

        # Cap target_k so we never ask for more communities than nodes.
        effective_k = max(2, min(target_k, max(2, N // 4)))
        algorithms = _make_algorithms(effective_k)

        for alg_name, algorithm in algorithms.items():
            # NAC1 with embeddings needs entities + embeddings passed through.
            extra_kwargs = {}
            if alg_name == "nac1" and effective_use_emb:
                extra_kwargs = {"entities": sub_entities, "embeddings": sub_embeddings}

            try:
                # NAC algorithms occasionally hang; gate them behind a thread with timeout.
                if alg_name in ("nac1", "nac2"):
                    container: Dict[str, Any] = {"ok": False, "elapsed": None, "err": None}

                    def _runner():
                        try:
                            _t0 = time.perf_counter()
                            algorithm.detect_communities(subgraph, **extra_kwargs)
                            container["elapsed"] = time.perf_counter() - _t0
                            container["ok"] = True
                        except Exception as e:
                            container["err"] = str(e)

                    th = threading.Thread(target=_runner)
                    th.start()
                    th.join(timeout=nac_timeout_seconds)
                    if th.is_alive():
                        elapsed = float("nan")
                        print(f"  {alg_name:>8}: TIMEOUT > {nac_timeout_seconds:.0f}s")
                    elif container["err"]:
                        elapsed = float("nan")
                        print(f"  {alg_name:>8}: FAILED ({container['err'][:80]})")
                    else:
                        elapsed = container["elapsed"]
                        print(f"  {alg_name:>8}: {elapsed:7.3f}s")
                else:
                    _t0 = time.perf_counter()
                    algorithm.detect_communities(subgraph)
                    elapsed = time.perf_counter() - _t0
                    print(f"  {alg_name:>8}: {elapsed:7.3f}s")

                timings[alg_name][N] = elapsed
            except Exception as e:
                timings[alg_name][N] = float("nan")
                print(f"  {alg_name:>8}: FAILED ({str(e)[:80]})")

        # Drop large per-graph caches so memory does not grow with N.
        try:
            clear_graphrag_caches()
        except Exception:
            pass

    # Plot wall-clock vs N (log-log to expose scaling exponents).
    if save_plot:
        try:
            label_map = {"c0": "Leiden", "spectral": "Spectral", "score": "SCORE", "nac1": "NAC1", "nac2": "NAC2"}
            fig, axes = plt.subplots(1, 2, figsize=(14, 6))
            for alg_name in ["c0", "spectral", "score", "nac1", "nac2"]:
                data = sorted(timings[alg_name].items())
                xs = [n for n, t in data if t == t]   # filter NaN
                ys = [t for n, t in data if t == t]
                if xs:
                    axes[0].plot(xs, ys, marker='o', label=label_map[alg_name], linewidth=2)
                    axes[1].plot(xs, ys, marker='o', label=label_map[alg_name], linewidth=2)
            axes[0].set_xlabel("Graph size N (entities)", fontsize=14)
            axes[0].set_ylabel("Community-detection wall-clock (s)", fontsize=14)
            axes[0].set_title("Linear scale", fontsize=16, fontweight='bold')
            axes[0].grid(True, alpha=0.3)
            axes[0].legend(fontsize=12)

            axes[1].set_xscale("log")
            axes[1].set_yscale("log")
            axes[1].set_xlabel("Graph size N (entities, log)", fontsize=14)
            axes[1].set_ylabel("Wall-clock (s, log)", fontsize=14)
            axes[1].set_title("Log–log (slope ≈ scaling exponent)", fontsize=16, fontweight='bold')
            axes[1].grid(True, which='both', alpha=0.3)
            axes[1].legend(fontsize=12)

            fig.suptitle("Scalability of community-detection methods",
                         fontsize=18, fontweight='bold')
            plt.tight_layout()
            plt.savefig(plot_filename, dpi=400, bbox_inches='tight')
            plt.close(fig)
            print(f"\nSaved scalability curve: {plot_filename}")
        except Exception as e:
            print(f"Could not save scalability plot: {e}")

    if save_csv:
        try:
            import csv
            with open(csv_filename, "w", newline="", encoding="utf-8") as f:
                writer = csv.writer(f)
                writer.writerow(["algorithm", "N", "seconds"])
                for alg_name, by_n in timings.items():
                    for N, t in sorted(by_n.items()):
                        writer.writerow([alg_name, N, f"{t:.6f}" if t == t else ""])
            print(f"Saved per-(algorithm, N) timings: {csv_filename}")
        except Exception as e:
            print(f"Could not save {csv_filename}: {e}")

    return {"sizes": node_sizes, "timings": timings, "target_k": target_k}


def save_timings_to_word(
    all_detection_timings: Dict[int, Dict[str, float]],
    graph_build_seconds: float,
    graph_stats: Dict[str, Any],
    cat_to_label: Dict[str, str],
    doc: Document = None,
) -> Document:
    """R1(a): write the per-algorithm community-detection wall-clock table into the docx."""
    if not DOCX_AVAILABLE:
        return doc
    if doc is None:
        doc = Document()

    doc.add_heading('Computational Cost (R1: Scalability)', level=1)
    if graph_stats:
        doc.add_paragraph(
            f"Graph used for timing: {graph_stats.get('n_nodes', 'N/A')} entities, "
            f"{graph_stats.get('n_edges', 'N/A')} relationships, "
            f"{graph_stats.get('n_components', 'N/A')} connected components, "
            f"density {graph_stats.get('density', 0):.4f}."
        )
    doc.add_paragraph(
        f"One-time graph construction (entity + relationship extraction, Leiden index): "
        f"{graph_build_seconds:.2f} seconds."
    )
    doc.add_paragraph(
        "Community-detection wall-clock per algorithm at each fixed k (seconds). "
        "Reported on the same shared graph above; all other pipeline costs (LLM "
        "summarisation, judging) are independent of the community-detection algorithm."
    )

    if not all_detection_timings:
        doc.add_paragraph("No timing data collected.")
        return doc

    methods_order = ["c0", "spectral", "score", "nac1", "nac2"]
    k_values = sorted(all_detection_timings.keys())

    table = doc.add_table(rows=1, cols=len(k_values) + 1)
    table.style = 'Light Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Algorithm'
    for i, k in enumerate(k_values):
        hdr_cells[i + 1].text = f'k = {k}'
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    for method in methods_order:
        # Only render rows we actually have data for.
        if not any(method in all_detection_timings[k] for k in k_values):
            continue
        row_cells = table.add_row().cells
        row_cells[0].text = cat_to_label.get(method, method)
        for i, k in enumerate(k_values):
            t = all_detection_timings[k].get(method)
            row_cells[i + 1].text = f"{t:.3f}" if t is not None else "N/A"

    doc.add_paragraph()
    doc.add_paragraph("Notes:")
    doc.add_paragraph("• Wall-clock measured with time.perf_counter() inside the experiment loop.", style='List Bullet')
    doc.add_paragraph("• Community-detection cost is reported in isolation; downstream LLM stages are algorithm-agnostic.", style='List Bullet')
    doc.add_paragraph("• See the scalability curve (scalability_curve.png) for behavior as graph size grows.", style='List Bullet')

    doc.add_page_break()
    return doc


def print_elo_summary(elo_ratings: Dict[str, Dict[str, float]], cat_to_label: Dict[str, str]) -> None:
    """Print Elo ratings summary table"""
    
    if not elo_ratings:

        return
    
    metrics = list(elo_ratings.keys())
    methods = list(next(iter(elo_ratings.values())).keys())
    
    avg_elo = {}
    for method in methods:
        ratings = [elo_ratings[metric][method] for metric in metrics if method in elo_ratings[metric]]
        avg_elo[method] = sum(ratings) / len(ratings) if ratings else 0
    sorted_methods = sorted(methods, key=lambda m: avg_elo[m], reverse=True)

    print("\n" + "="*80)
    print("ELO RATINGS SUMMARY")
    print("="*80)

    header = "Method".ljust(15)
    for metric in metrics:
        header += metric[:12].rjust(15)
    header += "Average".rjust(15)
    print(header)
    print("-" * len(header))

    for method in sorted_methods:
        row = cat_to_label.get(method, method).ljust(15)
        for metric in metrics:
            if metric in elo_ratings and method in elo_ratings[metric]:
                row += f"{elo_ratings[metric][method]:.0f}".rjust(15)
            else:
                row += "N/A".rjust(15)
        row += f"{avg_elo[method]:.0f}".rjust(15)
        print(row)
    print("\nInterpretation:")
    print("- Higher Elo rating = better performance")
    print("- Ratings are relative within each metric")
    print("- Average provides overall ranking across all metrics")


def create_winrate_matrix_graphs_for_k(winrate_results: Dict[str, Dict[str, Any]], k_value: int, save_plots: bool = True) -> None:
    """Create visualization heatmaps for win-rate matrices using a 2x2 layout
    with large fonts for print readability.
    """

    metrics = ["Comprehensiveness", "Diversity", "Empowerment", "Directness"]
    plt.rcParams.update({'font.size': 16})

    for k, k_results in winrate_results.items():
        if not any(metric in k_results for metric in metrics):
            continue

        # 2x2 grid so each heatmap is large enough to read in print.
        fig, axes = plt.subplots(2, 2, figsize=(14, 14))
        labels = k_results.get("labels", [])
        clean_labels = [label.split(' (k=')[0] if ' (k=' in label else label for label in labels]

        for idx, metric in enumerate(metrics):
            row = idx // 2
            col = idx % 2
            ax = axes[row, col]

            if metric in k_results:
                matrix = np.array(k_results[metric])

                ax.imshow(matrix, cmap='viridis', aspect='equal', vmin=0, vmax=1)

                # Annotate every cell with the win-rate percentage as an integer.
                # int() truncates toward zero (e.g. 66.6 -> 66), which dropped every
                # mirror pair (i,j)+(j,i) to 99 instead of 100. To guarantee each pair
                # sums to exactly 100, round the upper triangle and force every
                # lower-triangle cell to 100 minus its mirror (round() alone can still
                # land on 99/101 at exact .5 boundaries due to banker's rounding).
                for i in range(matrix.shape[0]):
                    for j in range(matrix.shape[1]):
                        if j >= i:
                            val = int(round(matrix[i, j] * 100))
                        else:
                            val = 100 - int(round(matrix[j, i] * 100))
                        ax.text(j, i, f'{val}', ha="center", va="center",
                                color="white" if matrix[i, j] < 0.5 else "black",
                                fontweight='bold', fontsize=22)

                ax.set_xticks(range(len(clean_labels)))
                ax.set_yticks(range(len(clean_labels)))
                ax.set_xticklabels(clean_labels, rotation=45, ha='right', fontsize=18)
                ax.set_yticklabels(clean_labels, fontsize=18)

                # Show the y-axis label on every left-column subplot in the 2x2 grid.
                if col == 0:
                    ax.set_ylabel('Win-Rate', fontweight='bold', fontsize=20)

                ax.set_title(metric, fontweight='bold', fontsize=24, pad=10)
            else:
                ax.text(0.5, 0.5, f'No data\nfor {metric}', ha='center', va='center',
                        transform=ax.transAxes, fontsize=18)
                ax.set_title(metric, fontweight='bold', fontsize=22)

        # Slightly more headroom between subplots to keep tick labels from overlapping titles.
        plt.tight_layout()
        plt.subplots_adjust(hspace=0.45, wspace=0.25)

        if save_plots:
            filename = os.path.join("Output", f"winrate_matrices_k_{k_value}.png")
            plt.savefig(filename, dpi=400, bbox_inches='tight')
            print(f"\nSaved winrate matrix visualization: {filename}")
            plt.close(fig)
        else:
            plt.show()

    plt.rcParams.update({'font.size': 10})

# Entry point
def main():
    """Main function implementing Microsoft's win-rate matrix experimental framework"""

    global global_leiden_graphrag
    global_leiden_graphrag = None
    
    folder_path = "raw_data"
    
    folder = Path(folder_path)
    if not folder.exists():
        folder_path = "."  

    fixed_doc_count = 70
    k_values_to_test = [10, 30, 50, 80, 100]
    OUTPUT_DIR = "Output"
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    
    print("\n" + "="*80)
    print("FIXED K EXPERIMENT - WIN-RATE ANALYSIS WITH MULTIPLE K VALUES")
    print("="*80)
    print(f"Loading {fixed_doc_count} documents...")
    documents, titles = load_multiple_documents(folder_path=folder_path, max_documents=fixed_doc_count)
    print(f"Loaded {len(documents)} documents")
    
    print("\nGenerating queries...")
    try:
        queries = generate_microsoft_algorithm1_questions(
            documents, 
            titles,
            K=4,  # User personas
            N=5,  # Tasks per persona  
            M=5   # Questions per task
        )
        print(f"Generated {len(queries)} queries")
    except Exception as e:
        print(f"Error generating queries: {e}")
        queries = []
    
    if not queries:
        print("No queries generated, exiting")
        return
    
    print("\nBuilding entity graph from documents...")
    try:
        temp_leiden = LeidenCommunityDetection()
        baseline_gr = MicrosoftGraphRAG(temp_leiden)
        # R1(a): time the entity/relationship extraction + graph build (one-time cost).
        _t_build_start = time.perf_counter()
        baseline_gr.index(documents, titles, max_cluster_size=50)
        graph_build_seconds = time.perf_counter() - _t_build_start
        print(f"Graph built successfully with {baseline_gr.graph.number_of_nodes()} nodes and {baseline_gr.graph.number_of_edges()} edges (build time: {graph_build_seconds:.2f}s)")
    except Exception as e:
        print(f"Failed to build graph: {e}")
        return

    all_winrate_results = {}
    all_elo_results = {}
    # R1(a): per-k per-algorithm community-detection timings.
    all_detection_timings: Dict[int, Dict[str, float]] = {}

    # Hoisted so post-loop code (timing/ELO summaries) is safe even if the
    # k-loop never executes or all algorithms fail in one iteration.
    cat_to_label = {
        "c0": "Leiden",
        "spectral": "Spectral",
        "score": "SCORE",
        "nac1": "NAC1",
        "nac2": "NAC2",
    }
    
    word_doc = None
    if DOCX_AVAILABLE:
        word_doc = Document()
        word_doc.add_heading('GraphRAG Community Detection ELO Ratings', 0)
        word_doc.add_paragraph(f'Experiment Date: {time.strftime("%Y-%m-%d %H:%M:%S")}')
        word_doc.add_paragraph(f'Number of Documents: {len(documents)}')
        word_doc.add_paragraph(f'Number of Queries: {len(queries)}')
        word_doc.add_paragraph(f'K Values Tested: {", ".join(map(str, k_values_to_test))}')
        word_doc.add_page_break()
    
    for k_value in k_values_to_test:
        
        winrate_results = run_winrate_matrix_experiment_fixed_k(
            documents, titles, queries, k_value, verbose=True, baseline_gr=baseline_gr
        )
        
        all_winrate_results[k_value] = winrate_results
        if "fixed_k" in winrate_results:
            result = winrate_results["fixed_k"]
            # R1(a): collect per-k detection timings.
            if "detection_timings" in result:
                all_detection_timings[k_value] = result["detection_timings"]

        available_methods = []
        labels = result.get("labels", [])

        # Build the list of methods present in this run, in the order their labels appear.
        for label in labels:
            if "Leiden" in label:
                available_methods.append("c0")
            elif "Spectral" in label:
                available_methods.append("spectral")
            elif "SCORE" in label:
                available_methods.append("score")
            elif "NAC1" in label:
                available_methods.append("nac1")
            elif "NAC2" in label:
                available_methods.append("nac2")

        # Compute ELO once per k after the method list is complete.
        elo_ratings = compute_elo_ratings(result, available_methods)
        all_elo_results[k_value] = elo_ratings

        print_elo_summary(elo_ratings, cat_to_label)
        
        if word_doc and DOCX_AVAILABLE:
            word_doc = save_elo_to_word(elo_ratings, cat_to_label, k_value, word_doc)
        
        print(f"\nCreating winrate matrix visualization for K={k_value}...")
        create_winrate_matrix_graphs_for_k(winrate_results, k_value, save_plots=True)
    
    _save_judge_cache(judge_cache)

    # R1(b): scalability curve via subgraph subsampling (no new LLM calls).
    scalability_results = {}
    try:
        scalability_results = run_scalability_experiment(
            baseline_gr=baseline_gr,
            node_sizes=None,        # auto, capped at the actual graph size
            target_k=50,            # fixed k for a fair comparison across N
            nac_timeout_seconds=120.0,
            save_plot=True,
            save_csv=True,
            plot_filename=os.path.join(OUTPUT_DIR, "scalability_curve.png"),
            csv_filename=os.path.join(OUTPUT_DIR, "scalability_timings.csv"),
        )
    except Exception as _e:
        print(f"Scalability experiment failed: {_e}")

    # R1(a): persist timings to JSON + write a timing table into the docx.
    timing_payload = {
        "graph_build_seconds": graph_build_seconds,
        "graph_stats": getattr(run_winrate_matrix_experiment_fixed_k, "_last_graph_stats", {}),
        "detection_timings_per_k": {str(k): v for k, v in all_detection_timings.items()},
    }
    try:
        _timing_path = os.path.join(OUTPUT_DIR, "timing_results.json")
        with open(_timing_path, "w", encoding="utf-8") as _f:
            json.dump(timing_payload, _f, indent=2)
        print(f"\nSaved per-algorithm timings to {_timing_path}")
    except Exception as _e:
        print(f"Could not save timing_results.json: {_e}")

    if word_doc and DOCX_AVAILABLE:
        word_doc = save_timings_to_word(
            all_detection_timings,
            graph_build_seconds,
            getattr(run_winrate_matrix_experiment_fixed_k, "_last_graph_stats", {}),
            cat_to_label,
            doc=word_doc,
        )

    if word_doc and DOCX_AVAILABLE:
        word_filename = os.path.join(OUTPUT_DIR, f"elo_ratings_summary_{time.strftime('%Y%m%d_%H%M%S')}.docx")
        word_doc.save(word_filename)
        print(f"\nSaved ELO ratings to: {word_filename}")

    # Final: produce baseline-relative ELO progression plots from all docx files in the folder.
    try:
        import subprocess, sys
        subprocess.run([sys.executable, "plot_elo_vs_baseline_from_docx.py"], check=False)
    except Exception as _e:
        print(f"Baseline ELO plot generation failed: {_e}")
    
    print("\n" + "="*80)
    print("EXPERIMENTS COMPLETED")
    print("="*80)
    print("\nGenerated files:")
    print("  - winrate_matrices_k_10.png: Win-rate heatmaps for K=10")
    print("  - winrate_matrices_k_50.png: Win-rate heatmaps for K=50")
    print("  - winrate_matrices_k_100.png: Win-rate heatmaps for K=100")
    print("  - elo_progression_by_k.png: Per-metric ELO vs k=10/50/100 (R2.M2)")
    print("  - elo_progression_avg_by_k.png: Average ELO vs k=10/50/100 (R2.M2)")
    print("  - scalability_curve.png: Wall-clock vs graph size N (R1.b)")
    print("  - scalability_timings.csv: Per-(algorithm, N) raw timings (R1.b)")
    print("  - timing_results.json: Per-(algorithm, k) wall-clock at the reported scale (R1.a)")
    if word_doc and DOCX_AVAILABLE:
        print(f"  - {word_filename}: ELO ratings + timing summary in Word format")
    print("\nThree ELO tables have been printed above for K=10, K=50, and K=100.")


if __name__ == "__main__":
    if not os.getenv("OPENAI_API_KEY"):
        print("Set OPENAI_API_KEY environment variable first!")
        exit(1)

    main()
